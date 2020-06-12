# !/usr/bin/env python
# encoding: utf-8

import os
import requests
import re
import hashlib
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from pathlib import Path


def getImage(image_url, md5_title, number):
    image_request = requests.get(image_url)
    image_name = md5_title + "_" + str(number) + ".jpg"
    open("images/" + image_name, 'wb').write(image_request.content)

    return image_name


def createDoc(title, content):
    document = Document()
    document.add_heading(title, 0)

    paragraph = document.add_paragraph(u'')
    md5 = hashlib.md5()
    md5.update(title.encode("utf-8"))
    md5_title = md5.hexdigest()
    image_number = 0
    for item in content:
        re_obj = re.compile(r'http[s]?://.*?(?:gif|png|jpg|jpeg|webp|svg|psd|bmp|tif)')
        url_match = re.findall(re_obj, item)
        if len(url_match) > 0:
            image_url = url_match[0]
            image_name = getImage(image_url, md5_title, image_number)
            document.add_picture("images/" + image_name, width=Inches(2.27))
            image_number += 1
        else:
            add_row = paragraph.add_run(item + "\n")
            add_row.font.size = Pt(12)
            add_row.font.name = u'微软雅黑'
            r = add_row._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    document.save("doc/" + title + ".docx")
    save_doc = Path("doc/" + title + ".docx")
    print("保存成功") if save_doc else print("保存失败")


def getPost(post_url):
    os.environ['NO_PROXY'] = post_url
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) " +
                      "Chrome/83.0.4103.97 Safari/537.36"
    }
    requests_obj = requests.get(post_url, headers=headers)
    bs_obj = BeautifulSoup(requests_obj.text, "html5lib")
    post_html = bs_obj.select_one("div.content")
    post_title = post_html.find("h2").text
    post_content = post_html.find_all("p")
    post_content_text = []
    for item in post_content:
        item_text = item.text
        item_text = item_text.strip().replace(u'\u3000', u' ').replace(u'\xa0', u' ')
        if len(item_text) > 0:
            post_content_text.append(item_text)
        else:
            img_element = item.img
            if img_element:
                post_content_text.append(img_element["src"])

    createDoc(post_title, post_content_text)


def oneUrl():
    url = input("please enter url: ")
    re_obj = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    url_match = re.findall(re_obj, url)
    if url_match:
        getPost(url_match[0])
    else:
        print("error url, please re-enter")
        oneUrl()


def index():
    app_mode = input("1. one url\n2. url list\nplease enter run mode: ")
    if app_mode == "1":
        oneUrl()
    elif app_mode == "2":
        file = open("url_list.txt")
        file_line_text = file.readline()
        line_number = 0
        while file_line_text:
            file_line_text = file_line_text.strip('\n')
            print(line_number, end=": ")
            print(file_line_text)
            getPost(file_line_text)
            file_line_text = file.readline()
            line_number += 1

        file.close()
    else:
        print("run mode error, re-enter")

    index()


if __name__ == '__main__':
    index()
